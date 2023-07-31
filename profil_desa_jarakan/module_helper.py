
import re
import os
import subprocess
from datetime import date

from django.conf import settings

from docx import Document

DOCX_TEMPLATES = {
    'contoh': os.path.join(settings.BASE_DIR, 'docx_templates' ,'surat_contoh.docx'),
    'sktm': os.path.join(settings.BASE_DIR, 'docx_templates' ,'sktm.docx'),
    'gagal': os.path.join(settings.BASE_DIR, 'docx_templates' ,'gagal.docx'),
}

INDONESIA_MONTHS = "Januari Februari Maret April Mei Juni Juli Agustus September Oktober November Desember".split()

def replace_text_in_docx(template_name, replacement_data):
    file_path = DOCX_TEMPLATES[template_name]
    doc = Document(file_path)
    builtin_data_replacements = {
        'TANGGAL_HARI_INI': f'{date.today().day} - {INDONESIA_MONTHS[date.today().month-1]} - {date.today().year}'
    }
    replacement_data.update(builtin_data_replacements)

    for key,value in replacement_data.items():
        value=str(value)
        for paragraph in doc.paragraphs:
            formatted_code = '{{%s}}'%(key)
            if formatted_code in paragraph.text:
                print('found', formatted_code)
                paragraph.text = paragraph.text.replace(formatted_code, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    formatted_code = '{{%s}}'%(key)
                    if formatted_code in cell.text:
                        print('found', formatted_code)
                        cell.text = cell.text.replace(formatted_code, value)

    out_path = os.path.join(settings.BASE_DIR, 'temp', 'out_'+os.path.basename(file_path))
    doc.save(out_path)
    return out_path

def docx2pdf(doc_path, delete_doc=True):
    if os.path.isfile(doc_path):
        output_dir = os.path.join(settings.BASE_DIR, 'temp')
        output_path = os.path.join(output_dir, os.path.basename(doc_path).split('.')[0]+'.pdf')
        cmd_output = subprocess.check_output(f'libreoffice --headless --convert-to pdf --outdir \'{output_dir}\' \'{doc_path}\'', shell=True)
        if os.path.isfile(output_path):
            os.remove(doc_path)
            return output_path, cmd_output.decode()
        raise FileNotFoundError(f'Invalid pdf path: {output_path}')
    raise FileNotFoundError(f'Invalid doc path: {doc_path}')

if __name__ == "__main__":
    # Replace these with your file path, old text, and new text
    print(replace_text_in_docx('contoh', {'NAMA_MAHASISWA': 'KEVIN'}))